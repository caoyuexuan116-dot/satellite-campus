from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("outputs")
DOCX_PATH = OUT_DIR / "no_music_did_results_report.docx"
LOG_PATH = OUT_DIR / "drop_central_music_sensitivity.log"


REGRESSION_ROWS = [
    ["论文数主回归", "0.0525", "0.0266", "0.083", "2,741,807", "正向，10% 水平边际显著"],
    ["WoS Core 引用", "0.7746", "0.7003", "0.301", "2,741,807", "正向，但不显著"],
    ["All Databases 引用", "0.7669", "0.7358", "0.328", "2,741,807", "正向，但不显著"],
    ["存量作者论文数", "0.2876", "0.1740", "0.137", "1,006,595", "系数更大，但不显著"],
]


PERSON_ROWS = [
    ["东北农业大学", "10,141", "对照组"],
    ["东北林业大学", "8,308", "对照组"],
    ["东华大学", "13,311", "处理组，2003"],
    ["中国传媒大学", "2,532", "对照组"],
    ["中央财经大学", "2,061", "处理组，2009"],
    ["北京理工大学", "29,405", "处理组，2007"],
    ["华东理工大学", "14,227", "处理组，2007"],
    ["广西大学", "16,680", "对照组"],
    ["暨南大学", "22,544", "处理组，2014"],
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(header_cells[i], header, bold=True)
        set_cell_shading(header_cells[i], "EDEDED")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)


def build_docx() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("去除中央音乐学院后的 DID 结果汇总")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("WOS 多校区与教师个人产出小样本测试")
    run.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_heading("一、数据处理口径", level=1)
    add_bullets(
        doc,
        [
            "Python 读取学校 WOS Excel，按 Author Full Names 拆分作者。",
            "用 Addresses 中的作者-机构地址块识别本校作者，只保留地址匹配当前学校别名的作者。",
            "用 UT + person_id 去重，避免同一学校同一作者同一篇论文重复计数。",
            "聚合为 2000-2022 年个人-年平衡面板，再交给 Stata 估计多时点 DID。",
            "本报告过程中剔除中央音乐学院；原因是该校仅 25 个 person_id，样本量过小。",
        ],
    )

    doc.add_heading("二、样本检查", level=1)
    p = doc.add_paragraph()
    p.add_run("剔除中央音乐学院后，样本为 9 所学校、119,209 名作者、2,741,807 个个人-年观测。").bold = True
    doc.add_paragraph("处理组保持 5 所；对照组为中国传媒大学、东北林业大学、东北农业大学、广西大学。")
    add_table(doc, ["学校", "person_id 数量", "组别/处理时间"], PERSON_ROWS)

    doc.add_heading("三、多时点 DID 检查", level=1)
    add_bullets(
        doc,
        [
            "东华大学在 2003 年切换为 post=1。",
            "北京理工大学、华东理工大学在 2007 年切换为 post=1。",
            "中央财经大学在 2009 年切换为 post=1。",
            "暨南大学在 2014 年切换为 post=1。",
            "对照组 post 始终为 0，Stata 断言 did = treated × post 通过。",
        ],
    )

    doc.add_heading("四、回归结果", level=1)
    add_table(doc, ["模型", "DID 系数", "标准误", "p 值", "N", "解释"], REGRESSION_ROWS)

    doc.add_heading("五、结果解读", level=1)
    add_bullets(
        doc,
        [
            "论文数主回归系数为 0.0525，表示新校区启用后处理组作者平均每人每年发文约增加 0.053 篇；在 10% 水平边际显著。",
            "引用指标均为正，但 p 值较大，当前不能认为引用影响显著。",
            "存量作者样本的论文数系数为 0.2876，高于主回归，暗示产出提升可能更多体现在改革前已存在的作者身上。",
            "由于只有 9 所学校、学校聚类数量少，当前结果应定位为小样本管线测试，而非正式因果结论。",
        ],
    )

    doc.add_heading("六、注意事项", level=1)
    add_bullets(
        doc,
        [
            "虽然已用 Addresses 排除外校合作者，但姓名消歧仍未完全解决。",
            "对照组数量较少，后续应扩展学校样本并补充事件研究和平行趋势检验。",
            "引用变量存在发表后的滞后积累，短期 DID 解释需谨慎。",
        ],
    )

    doc.add_paragraph("对应 Stata 日志：outputs/drop_central_music_sensitivity.log")

    doc.add_page_break()
    doc.add_heading("附录：Stata 完整日志", level=1)
    doc.add_paragraph("以下为剔除中央音乐学院后，本次 Stata 运行产生的完整文本日志，便于逐项核对。")
    log_text = LOG_PATH.read_text(encoding="utf-8", errors="replace")
    for chunk_start in range(0, len(log_text), 2500):
        p = doc.add_paragraph()
        run = p.add_run(log_text[chunk_start : chunk_start + 2500])
        run.font.name = "Courier New"
        run.font.size = Pt(7)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
